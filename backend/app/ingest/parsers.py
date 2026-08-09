"""Document parsers + mock source adapters (ingestion-system.md §2, §4.1).

The MVP uses "mock integrations" (roadmap-mvp.md Week 2): adapters that emit the
same normalized shape as production providers (Zoom/Meet, Drive/Notion, Slack,
Gmail, Airtable). Parsing real provider payloads is a Phase 2 adapter swap.
"""

import hashlib
import re
from dataclasses import dataclass, field


@dataclass
class ParsedDocument:
    doc_type: str
    title: str
    content: str
    source: str
    category: str | None = None
    brands: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def sha256(self) -> str:
        return hashlib.sha256(self.content.encode("utf-8")).hexdigest()


def parse_plain_text(text: str) -> str:
    return text.strip()


def parse_markdown(text: str) -> str:
    text = re.sub(r"```(?:json|python|bash)?\n.*?\n```", "", text, flags=re.DOTALL)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    return text.strip()


def parse_docx(path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()


def parse_pdf(path: str) -> str:  # pragma: no cover — optional; needs pdfplumber
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber not installed; install it to parse PDFs")
    with pdfplumber.open(path) as pdf:
        return "\n".join((page.extract_text() or "") for page in pdf.pages).strip()


# ---------------------------------------------------------------------- #
# mock provider adapters (deterministic, for MVP pipeline + seeding)
# ---------------------------------------------------------------------- #
def mock_meeting(title: str, transcript: str, brands: list[str] | None = None, meeting_id: str | None = None) -> ParsedDocument:
    return ParsedDocument(
        doc_type="meeting",
        title=title,
        content=transcript,
        source="zoom",
        category=None,
        brands=brands or ["all"],
        metadata={"meeting_id": meeting_id or "", "adapter": "mock_zoom"},
    )


def mock_playbook(title: str, content: str, brands: list[str] | None = None, version: str = "1") -> ParsedDocument:
    return ParsedDocument(
        doc_type="playbook",
        title=title,
        content=content,
        source="drive",
        category=None,
        brands=brands or ["all"],
        metadata={"version": version, "adapter": "mock_drive"},
    )


def mock_decision(statement: str, rationale: str, category: str, outcome: str | None = None) -> ParsedDocument:
    body = f"Decision: {statement}\nRationale: {rationale}"
    if outcome:
        body += f"\nOutcome: {outcome}"
    return ParsedDocument(
        doc_type="decision",
        title=statement[:80],
        content=body,
        source="csv",
        category=category,
        brands=["all"],
        metadata={"outcome": outcome or ""},
    )


def mock_slack_digest(channel: str, thread_text: str) -> ParsedDocument:
    return ParsedDocument(
        doc_type="slack_digest",
        title=f"# {channel} thread",
        content=thread_text,
        source="slack",
        metadata={"channel": channel, "adapter": "mock_slack"},
    )
